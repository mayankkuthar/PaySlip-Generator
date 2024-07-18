from docx import Document
import pandas as pd
import os
import inflect
from datetime import datetime
from docx2pdf import convert
import smtplib
from email.message import EmailMessage

# Create an inflect engine
p = inflect.engine()

# Function to convert a number to words
def number_to_words(number):
    return p.number_to_words(number)

def replace_text_in_paragraph(paragraph, replacements):
    for run in paragraph.runs:
        for old_text, new_text in replacements.items():
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)

def replace_text_in_doc(input_file_path, output_file_path, replacements):
    # Load the document
    doc = Document(input_file_path)

    # Replace text in paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

    # Replace text in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

    # Save the modified document
    doc.save(output_file_path)
    convert(output_file_path)
    os.remove(output_file_path)
    print("Saved to "+os.path.splitext(output_file_path)[0] + '.pdf')

def mail_slip(subject, body, to_email, from_email, password, attachment_path):
    # Create the email message
    msg = EmailMessage()
    msg['Subject'] = subject
    msg['From'] = from_email
    msg['To'] = to_email
    msg.set_content(body)

    # Add the attachment
    with open(attachment_path, 'rb') as file:
        file_name = os.path.basename(attachment_path)
        msg.add_attachment(file.read(), maintype='application', subtype='octet-stream', filename=file_name)

    try:
        # Use SMTP_SSL for secure connection
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
            server.login(from_email, password)
            server.send_message(msg)
            print("Email sent successfully to "+ to_email)
    except smtplib.SMTPException as e:
        print(f"Failed to send email: {e}")
        

# Example usage:
input_file = 'Payslip.docx'

# Specify the path to your Excel file
file_path = 'Salary_slips_generator.xlsx'
# Read the sheet "Salary_slips" starting from row 2 (index 1 in zero-indexed)
df = pd.read_excel(file_path, sheet_name='Salary_slips', skiprows=1)

# Iterate through the DataFrame row by row
for index, row in df.iterrows():
    EID = row["EID"]
    Name = row['Name']
    Designation = row['Designation']
    Basic = row['Basic']
    HRA = row['HRA']
    Special_Allowance = row['Special Allowance']
    Conveyance_Allowance = row['Conveyance Allowance']
    Income_Tax = row['Income Tax']
    Provident_Fund = row['Provident Fund']
    Other_Allowance = row['Other Allowance']
    Net_Pay = row['Net Pay']
    Month = row['Month'].split("'")[0]
    Year = "20"+row['Month'].split("'")[1]
    Email = row['Email']
    Paid_Days = row['Paid Days']
    LOP = row['LOP']
    PF_No = row['PF No']

    # Create directory structure if it doesn't exist
    directory = os.path.join(Year, Month)
    if not os.path.exists(directory):
        os.makedirs(directory)

    file = f'{EID}_{Name}_{row["Month"]}_payslip.docx'
    output_file = os.path.join(directory, file)
    replacements = {
        'XmonthX': Month,
        'XyearX': Year,
        'XrupeeX': "₹{:,.0f}".format(Net_Pay),
        'XsX': "₹{:,.0f}".format(Basic),
        'XyX': "₹{:,.0f}".format(HRA),
        'XcaX': "₹{:,.0f}".format(Conveyance_Allowance),
        'XoaX': "₹{:,.0f}".format(Other_Allowance),
        'XsaX': "₹{:,.0f}".format(Special_Allowance),
        'XdysX': "₹{:,.0f}".format(Income_Tax),
        'XlpdaysX': "₹{:,.0f}".format(Provident_Fund),
        'XrwX': number_to_words(Net_Pay).capitalize()+" Only",
        'XenameX': Name,
        'XeidX': str(EID),
        'XdateX': datetime.today().strftime('%Y-%m-%d'),
        'XdaysX': str(Paid_Days),
        'XlopdaysX': str(LOP),
        'XpfnoX': PF_No,
        'XdesX': Designation,
    }
    replace_text_in_doc(input_file, output_file, replacements)

    if(row['Mail'] == "Yes"):
        subject = " Your Pay Slip || Emulus"
        body = "This is your monthly Pay Slip from Emulus."
        from_email = "mayank.emulus@gmail.com"
        password = "jdau qykc lzwb gpoa"
        attachment_path = os.path.splitext(output_file)[0] + '.pdf'
        mail_slip(subject, body, Email, from_email, password, attachment_path)
